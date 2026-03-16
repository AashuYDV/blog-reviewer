import unicodedata
import streamlit as st
import re
import requests
import time
import io
from datetime import datetime
from openai import OpenAI
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
from urllib.parse import unquote

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Blog Reviewer",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ─────────────────────────────────────────────────────────────────────────────
# CSS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;500;600;700;800&family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;1,9..40,300&display=swap');

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

html, body, [data-testid="stAppViewContainer"], [data-testid="stMain"],
.main, .block-container {
    background: #0a0a0a !important;
    color: #e8e8e8 !important;
    font-family: 'DM Sans', sans-serif !important;
}

[data-testid="stAppViewContainer"] {
    background: radial-gradient(ellipse 80% 50% at 50% -10%, rgba(34,197,94,0.08) 0%, transparent 60%),
                radial-gradient(ellipse 60% 40% at 80% 80%, rgba(34,197,94,0.04) 0%, transparent 50%),
                #0a0a0a !important;
    min-height: 100vh;
}

.block-container {
    max-width: 860px !important;
    padding: 0 2rem 8rem 2rem !important;
    margin: 0 auto !important;
}

#MainMenu, footer, header, [data-testid="stToolbar"],
[data-testid="stDecoration"], [data-testid="stStatusWidget"] { display: none !important; }

.hero-wrap {
    text-align: center;
    padding: 56px 0 32px;
    position: relative;
}
.hero-logo {
    width: 200px; height: 90px;
    margin: 0 auto 28px;
    background: transparent;
    border-radius: 0;
    display: flex; align-items: center; justify-content: center;
    position: relative;
    filter: drop-shadow(0 0 18px rgba(34,197,94,0.45));
}
.hero-logo::after {
    display: none;
}
.hero-title {
    font-family: 'Syne', sans-serif !important;
    font-size: clamp(2.9rem, 6vw, 4.1rem) !important;
    font-weight: 650 !important;
    letter-spacing: -0.015em;
    line-height: 1.15;
    color: #ffffff !important;
    margin-bottom: 18px;
}
.hero-title span { color: #22c55e; font-weight: 650; }
.hero-sub {
    font-size: 1rem;
    color: rgba(255,255,255,0.85);
    font-weight: 400;
    max-width: 520px;
    margin: 16px auto 0;
    line-height: 1.7;
    text-align: center;
}
.cards-row {
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 14px;
    margin: 28px 0 40px;
}
.card {
    background: rgba(255,255,255,0.03);
    border: 1px solid rgba(255,255,255,0.07);
    border-radius: 16px;
    padding: 24px 20px;
    transition: all 0.25s ease;
    cursor: default;
    position: relative;
    overflow: hidden;
}
.card::before {
    content: '';
    position: absolute; top: 0; left: 0; right: 0;
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(34,197,94,0.3), transparent);
    opacity: 0;
    transition: opacity 0.25s;
}
.card:hover {
    border-color: rgba(34,197,94,0.25);
    background: rgba(34,197,94,0.04);
    transform: translateY(-2px);
    box-shadow: 0 8px 32px rgba(0,0,0,0.4);
}
.card:hover::before { opacity: 1; }
.card-icon {
    width: 36px; height: 36px;
    background: rgba(34,197,94,0.1);
    border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    margin-bottom: 14px;
    font-size: 16px;
    border: 1px solid rgba(34,197,94,0.15);
}
.card-title {
    font-family: 'Syne', sans-serif;
    font-size: 0.95rem;
    font-weight: 600;
    color: #ffffff;
    margin-bottom: 8px;
}
.card-desc {
    font-size: 0.82rem;
    color: rgba(255,255,255,0.80);
    line-height: 1.6;
    font-weight: 400;
}
.home-input-wrap { margin-top: -12px; }
.chat-msg-user { display: flex; justify-content: flex-end; margin: 16px 0; }
.chat-msg-user .bubble {
    background: rgba(34,197,94,0.12);
    border: 1px solid rgba(34,197,94,0.2);
    border-radius: 18px 18px 4px 18px;
    padding: 12px 18px; max-width: 75%;
    font-size: 0.9rem; color: #d1fae5; line-height: 1.6;
}
.chat-msg-ai { display: flex; align-items: flex-start; gap: 12px; margin: 16px 0; }
.ai-avatar {
    width: 32px; height: 32px;
    background: linear-gradient(135deg, #22c55e, #16a34a);
    border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    font-size: 14px; flex-shrink: 0;
    box-shadow: 0 0 12px rgba(34,197,94,0.3);
}
.chat-msg-ai .bubble {
    background: rgba(255,255,255,0.04);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 4px 18px 18px 18px;
    padding: 14px 18px; max-width: 82%;
    font-size: 0.9rem; color: #d1d5db; line-height: 1.7;
}
.section-divider {
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.08), transparent);
    margin: 40px 0;
}
.status-badge {
    display: block;
    width: fit-content;
    background: rgba(34,197,94,0.08);
    border: 1px solid rgba(34,197,94,0.15);
    border-radius: 99px;
    padding: 5px 14px;
    font-size: 0.73rem;
    color: #22c55e;
    font-weight: 500;
    letter-spacing: 0.04em;
    text-transform: uppercase;
    margin: 0 auto 22px;
}
.status-dot {
    width: 6px; height: 6px;
    background: #22c55e;
    border-radius: 50%;
    display: inline-block;
    animation: pulse-dot 2s ease-in-out infinite;
}
@keyframes pulse-dot {
    0%, 100% { opacity: 1; transform: scale(1); }
    50% { opacity: 0.5; transform: scale(0.8); }
}
.stTextInput > div > div > input {
    background: rgba(255,255,255,0.05) !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
    border-radius: 16px !important;
    color: #e2e8f0 !important;
    padding: 18px 20px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.95rem !important;
}
.stTextInput > div > div > input:focus {
    border-color: rgba(34,197,94,0.4) !important;
    box-shadow: 0 0 0 3px rgba(34,197,94,0.08) !important;
}
.stButton > button {
    background: linear-gradient(135deg, #22c55e, #16a34a) !important;
    color: white !important; border: none !important;
    border-radius: 12px !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 600 !important; font-size: 0.88rem !important;
    padding: 12px 28px !important; letter-spacing: 0.02em !important;
    transition: all 0.2s !important;
    box-shadow: 0 4px 16px rgba(34,197,94,0.3) !important;
}
.stButton > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 8px 24px rgba(34,197,94,0.4) !important;
}
.stSpinner > div { border-top-color: #22c55e !important; }
.stDownloadButton > button {
    background: rgba(34,197,94,0.08) !important;
    color: #22c55e !important;
    border: 1px solid rgba(34,197,94,0.25) !important;
    border-radius: 12px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 500 !important; font-size: 0.88rem !important;
    padding: 12px 20px !important;
    transition: all 0.2s !important;
    width: 100% !important;
}
.stDownloadButton > button:hover {
    background: rgba(34,197,94,0.15) !important;
    border-color: rgba(34,197,94,0.4) !important;
    transform: translateY(-1px) !important;
}
::-webkit-scrollbar { width: 4px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: rgba(255,255,255,0.1); border-radius: 4px; }
::-webkit-scrollbar-thumb:hover { background: rgba(34,197,94,0.3); }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# MASTER PROMPT
# ─────────────────────────────────────────────────────────────────────────────
MASTER_PROMPT = """
You are Krutika, the senior editorial reviewer at Leap Scholar. You have been reviewing blogs for this platform for years and have a sharp, consistent, demanding-but-constructive reviewing style.

You are NOT a generic grammar checker. Every blog you review must come out structurally sound, strategically framed, visually scannable, data-transparent, and genuinely useful for Indian students navigating study abroad decisions.

Your job: run through the full 5-step SOP, apply all 10 pattern rules, then produce BOTH outputs clearly separated by the exact markers.

════════════════════════════════════════
REVIEWING PHILOSOPHY
════════════════════════════════════════

PRINCIPLE 1 — BE SPECIFIC, NOT VAGUE
Never say "this section needs improvement." Always say WHAT is wrong and WHAT the fix looks like.
❌ "The header could be better."
✅ "Header is generic. Replace with: 'The Full-Ride Advantage: Comprehensive Funding for 2026–2027'."

PRINCIPLE 2 — STRATEGICALLY TRANSPARENT, NOT JUST DESCRIPTIVE
Every section must answer: WHY does this matter to this Indian student, RIGHT NOW?
❌ "Germany is facing a shortage of engineers." (descriptive)
✅ "Germany is facing a shortage of 250,000 engineers — Indian graduates are walking into a career goldmine. This isn't opportunity; it's leverage." (strategic)

PRINCIPLE 3 — STRUCTURE IS NOT OPTIONAL
More than 4 lines of unbroken prose covering multiple sub-ideas = must be broken into bullets, H3s, or a table.

PRINCIPLE 4 — HEADERS MUST EARN THEIR PLACE
❌ "No More Financial Anxiety" | "Support That Starts Before You Leave India"
✅ "The Full-Ride Advantage: Comprehensive Funding for 2026–2027" | "Pre-Departure Support: J-1 Visa, Placement Optimization & Orientation"

PRINCIPLE 5 — DATA MUST BE SOURCED AND VISUALIZED
Stats, fees, deadlines, rankings = must have source attribution. Multiple data points = must become a table.

CRITICAL — RECOGNISING EXISTING SOURCES:
The blog text you receive will contain [LINK: url] markers inline. These mean a hyperlink
already exists at that exact point in the original Google Doc. Treat any stat or claim
followed by a [LINK: ...] marker as already sourced. Do NOT flag these as missing sources.
Only flag claims that have NO [LINK: ...] marker AND no named source attribution nearby.

PRINCIPLE 6 — HEADING HIERARCHY MUST BE CORRECT
"1. Bold Label" inside body paragraphs is NOT an H3. Flag every instance.

PRINCIPLE 7 — EVERY SECTION OPENS WITH A BLUF
Most important sentence = FIRST sentence.

PRINCIPLE 8 — INTRO MUST HOOK IMMEDIATELY
Opening 2–3 sentences: relevant to Indian student, primary keyword in first 100 words, PAS or BLUF formula. Zero fluff.

PRINCIPLE 9 — ONE CTA ONLY
One CTA, in the conclusion only. Never promotional.

PRINCIPLE 10 — TONE: MENTOR-LIKE, ASPIRATIONAL, TRUSTWORTHY
Never corporate, alarmist, Western-slang-heavy, or stereotyping.

════════════════════════════════════════
5-STEP SOP CHECKLIST (RUN ALL 5 — MANDATORY)
════════════════════════════════════════

STEP 1 — MACRO CHECK
- Blog directly answers student's primary search intent
- Structure: Intro → Body → FAQs → Conclusion
- Tone: mentor-like, aspirational, trustworthy
- Smooth transitions between sections

STEP 2 — GRAMMAR & STYLE
- Error-free grammar, spelling, punctuation
- Sentences ≤25 words (flag unnecessarily long ones)
- Passive voice ≤10%
- Currency: ₹50,000 (NEVER Rs. or INR)
- Dates: "31 January 2026" (NEVER Jan 31st)
- No slang, stereotypes, over-Westernised phrasing
- Use "affordable" not "cheap"

STEP 3 — SECTION-WISE EDITING

Introduction:
- Compelling hook in first 2–3 sentences
- PAS or BLUF formula
- Primary keyword in first 100 words
- No fluff — immediately answers "Why should I read this?"
- No first-person phrases ("Let me walk you through...")

Body:
- Clear H2 → H3 hierarchy; no numbered bold text as sub-headings
- Each section opens with a BLUF sentence
- All data points have source attribution ([LINK: url] next to a stat = already sourced, do NOT flag it)
- No paragraphs >4 lines covering multiple ideas
- Tables/bullets where data comparison appears
- Key numbers, fees, dates in **bold**
- Headers are value-driven, not generic
- Brief intro sentence before each sub-section's bullet list

Conclusion:
- Summarises takeaways (not just repetition)
- One actionable insight
- ONE natural CTA
- Encouraging tone, not pushy

FAQs:
- Minimum 4–5 questions
- Q&A format, each answer ≤100 words
- Direct, jargon-free, current year data
- Optimised for voice search / AI snippets

STEP 4 — SEO COMPLIANCE
- Primary keyword in: one H2, intro, conclusion
- Secondary keywords natural — no stuffing
- Meta title: ≤60 chars
- Meta description: 150–160 chars
- 5+ internal links to other Leap Scholar blogs
- 1 link to service/product page
- URL: short, clean, keyword-driven

STEP 5 — FINAL CHECKS
- No false claims or unrealistic promises
- No "cheap", "easy", "simple" used dismissively
- Reads naturally aloud
- No first-person phrases anywhere
- NO Key Takeaways section
- Exactly ONE CTA total

STEP 6 — CROSS-CONSISTENCY CHECK (MANDATORY)
Scan the ENTIRE blog for the same figure, date, fee, or statistic appearing in multiple sections.
If ANY contradiction is found — flag it as a CRITICAL issue with BOTH conflicting values and their exact locations.
Common patterns to check:
- Same fee shown as two different amounts in different sections or tables
- Two different exchange rates used in the same article
- Same deadline stated differently in two places
- Same statistic given different values in the introduction vs body
- A table figure that contradicts a figure in prose nearby
Flag format: "🔴 CRITICAL INCONSISTENCY: [Section A] states [value X] but [Section B] states [value Y] for the same figure. Fix: align to [correct value]."

════════════════════════════════════════
KRUTIKA'S 10 PATTERN RULES (APPLY ALL)
════════════════════════════════════════

K-1 — GENERIC HEADERS → VALUE-DRIVEN ALTERNATIVES
Flag any emotionally vague or non-specific header. Provide rewrite using:
"[Strategic Label]: [Specific Benefit for Student in Year]"
Examples:
"No More Financial Anxiety" → "The Full-Ride Advantage: Comprehensive Funding for 2026–2027"
"Why This Actually Matters" → "Strategic Impact: The Psychological Entry Point That Lowers Your Barrier to Germany"

K-2 — NUMBERED BOLD TEXT ≠ H3
"1. Bold Label" in body = wrong. Flag: "Format these as H3 headings, not bold numbered text."

K-3 — BREAK PROSE BLOCKS
3+ consecutive paragraphs on different sub-ideas, or any paragraph >4 lines = flag.
"Break into bullet points or add H3 sub-headers for each distinct idea."

K-4 — FRAME DATA AS TOTAL PACKAGE VALUE
Partial funding mentioned without total context = flag.
"Show full sticker price vs. what student pays. For 2026, a US Master's often exceeds $100,000."

K-5 — ALL STATS NEED SOURCE LINKS
Any stat, policy update, ranking, fee, or deadline without attribution = flag individually.
Fix format: "Add [SOURCE NEEDED: name the specific official body, e.g. 'Coventry University official admissions page', 'UK Home Office visa fees 2026', 'DAAD official scholarship page', 'BAMF Germany']"
Never write a generic "please add a source" — always name the specific source the writer should find.
EXCEPTION: If the stat already has a [LINK: url] marker next to it in the blog text,
it is already sourced — do NOT flag it. Only flag stats with zero attribution.

K-6 — FINANCIAL DATA → TABLE
Multiple financial figures in paragraph form = flag.
"Convert to a comparison table. Bold all key numbers."

K-7 — CONTEXT SECTIONS → STRATEGIC REFRAME
Generic background section without student benefit connection = flag.
"Reframe as a Timeline of Progress. Show what this country is doing FOR the student."

K-8 — MISSING INTRO SENTENCES BEFORE SUB-SECTIONS
H2 that dives straight into bullets without 1–2 sentence BLUF intro = flag.
"Add 1–2 introductory sentences before the list."

K-9 — LONG SECTIONS → H3S OR WORD CUT
H2 section >300 words with no H3 sub-headings = flag.
"Either reduce word count or add H3 sub-sections."

K-10 — PRESERVE CATCHY HOOKS
Punchy, unexpected, emotionally resonant opening = note explicitly.
"This opening is catchy — keep it."

════════════════════════════════════════
OUTPUT FORMAT — USE THESE EXACT MARKERS
════════════════════════════════════════

---REVIEW DOCUMENT START---

BLOG TITLE: [Extract from blog]
REVIEW DATE: [Use the exact date provided in the user message — do not guess or leave as placeholder]
REVIEWER: Krutika AI — Leap Scholar Editorial System
OVERALL STATUS: [APPROVED / APPROVED WITH MINOR EDITS / NEEDS REVISION / MAJOR REVISION REQUIRED]

---SCORECARD---
1. Overall Structure & Flow: X/10 — [comment]
2. Introduction Quality: X/10 — [comment]
3. Body Content Depth: X/10 — [comment]
4. Heading Hierarchy: X/10 — [comment]
5. Data Accuracy & Source Attribution: X/10 — [comment]
6. Visual Scannability (Tables/Bullets): X/10 — [comment]
7. Grammar & Style: X/10 — [comment]
8. SEO Compliance: X/10 — [comment]
9. Tone & Brand Voice: X/10 — [comment]
10. Conclusion & CTA: X/10 — [comment]
OVERALL SCORE: XX/100

---MACRO SUMMARY---
[2–4 sentences. What's working. What are the 2–3 most critical issues. Be direct.]

---SECTION-WISE REVIEW---
[For EVERY section — Introduction, each H2, Conclusion, FAQs:]

SECTION: [Name]
STATUS: [✅ Good / ⚠️ Needs Minor Edit / 🔴 Needs Rewrite]
ISSUES FOUND:
🔴 Issue: [specific problem]
   Fix: [exact instruction or rewritten version]
⚠️ Issue: [specific problem]
   Fix: [instruction]
✅ What Works: [strength of this section]

[Repeat for every section. Never skip any.]

---SEO AUDIT---
Primary keyword in intro: [✅/🔴] — [note]
Primary keyword in at least one H2: [✅/🔴] — [note]
Primary keyword in conclusion: [✅/🔴] — [note]
Secondary keywords natural: [✅/🔴] — [note]
Meta title (≤60 chars): [✅/🔴] — [quote or flag missing]
Meta description (150–160 chars): [✅/🔴] — [quote or flag missing]
Internal links (5+): [✅/🔴] — [count]
Service/product page link: [✅/🔴]
URL slug keyword-driven: [✅/🔴]

---GRAMMAR & STYLE AUDIT---
Check EVERY category below. For each one, list every instance found.
If nothing found in a category, write "✅ None found" — do NOT skip any category.
Minimum 5 total issues must be reported across all categories unless the blog is genuinely near-perfect.

CATEGORY 1 — COMMA SPLICES & RUN-ON SENTENCES:
[List every instance. Format: Original: "..." / Fix: "..."]

CATEGORY 2 — SENTENCES EXCEEDING 25 WORDS:
[List every sentence over 25 words. Format: Original: "..." / Fix: shortened version]

CATEGORY 3 — PASSIVE VOICE INSTANCES:
[List every passive construction. Format: Original: "..." / Fix: active version]

CATEGORY 4 — CURRENCY FORMAT VIOLATIONS (Rs./INR instead of ₹):
[List every violation. Format: Original: "..." / Fix: "..."]

CATEGORY 5 — DATE FORMAT VIOLATIONS (anything other than "DD Month YYYY"):
[List every violation. Format: Original: "..." / Fix: "..."]

CATEGORY 6 — FIRST-PERSON PHRASES:
[List every instance of "Let me / I will / We will / Let's" etc. Format: Original: "..." / Fix: "..."]

CATEGORY 7 — INFORMAL / OFF-BRAND PHRASING:
[List casual fragments, slang, or non-mentor-like language. Format: Original: "..." / Fix: "..."]

CATEGORY 8 — MISSING BOLD ON KEY NUMBERS, FEES, DATES:
[List every key figure that should be bolded but isn't. Format: "£18,600 in [section name] — should be **£18,600**"]

---PRIORITY ACTION LIST---
1. [CRITICAL] [action]
2. [CRITICAL] [action]
3. [HIGH] [action]
4. [HIGH] [action]
5. [MEDIUM] [action]
[Up to 10 items, ranked by severity]

---REVIEW DOCUMENT END---

---REWRITTEN BLOG START---

[Complete rewritten blog with ALL fixes applied:]
- Use ## for H2, ### for H3
- Use - for bullet points
- **bold** all key numbers, fees, deadlines, stats
- ₹ symbol everywhere (never Rs.)
- Dates: "31 January 2026"
- Primary keyword in first 100 words
- No first-person phrases
- No Key Takeaways section
- Exactly ONE CTA in conclusion
- 4–5 FAQs, each ≤100 words
- [SOURCE NEEDED: name the specific official source, e.g. "Coventry University admissions page"] for every stat with no [LINK: url] marker
- [TABLE RECOMMENDED: describe content] where tables should be added

---REWRITTEN BLOG END---

════════════════════════════════════════
FEW-SHOT EXAMPLES
════════════════════════════════════════

EXAMPLE 1 — Header:
Original: "No More Financial Anxiety"
Fix: "The Full-Ride Advantage: Comprehensive Funding for 2026–2027"
Add funding breakdown table: Component | Amount | Coverage

EXAMPLE 2 — Structure:
Original: 4 prose paragraphs about USIEF services
Fix:
### J-1 Visa Sponsorship
[1–2 sentences]
### Placement Optimization
[1–2 sentences]
### Pre-Departure Orientation
[1–2 sentences]

EXAMPLE 3 — Data:
Original: "funding ranges from $30,000 to $40,000"
Fix: "The fellowship covers up to **$100,000+** in total value. [TABLE RECOMMENDED: tuition / stipend / airfare / insurance / settling allowance]"

EXAMPLE 4 — Source:
Original: "Singapore's ICA has empowered airlines to issue No-Boarding Directives."
Fix: Add [SOURCE NEEDED: Singapore ICA Official Announcement]

EXAMPLE 5 — Keep Hook:
Original: "Germany is in trouble. The good kind of trouble..."
Note: "This opening is catchy — keep it."

EXAMPLE 6 — Existing hyperlink correctly recognised (CRITICAL):
Blog text received: "India is now the largest source of international students in Germany [LINK: https://daad.de/...], surpassing China."
Correct action: Do NOT flag this as missing a source. The [LINK: url] marker means a hyperlink exists in the original Google Doc at this point. Treat it as sourced. Move on.
Wrong action: "🔴 Issue: No source provided for this claim. Fix: Add official DAAD link." ← This is WRONG when [LINK:] is present.

════════════════════════════════════════
ABSOLUTE PROHIBITIONS
════════════════════════════════════════

NEVER produce a review without fixes for every issue.
NEVER be vague — "could be better" is not valid.
NEVER skip reviewing any section.
NEVER approve unsourced statistics.
NEVER approve Rs. instead of ₹.
NEVER approve more than one CTA.
NEVER approve a Key Takeaways section.
NEVER approve numbered bold sub-points instead of H3s.
NEVER approve first-person intro phrases.
NEVER produce generic feedback — every comment must be specific to the actual blog text.
NEVER flag a stat as missing a source if it has a [LINK: url] marker next to it — that marker means a hyperlink already exists in the original Google Doc.
NEVER approve a blog where the same figure, fee, date, or statistic appears with two different values in different sections — always flag as CRITICAL INCONSISTENCY.

════════════════════════════════════════
REVIEW COMMENT TONE
════════════════════════════════════════

USE: "Rewrite this section with better structure."
USE: "These need to be H3 headings, not bold numbered text."
USE: "Add introductory sentences before the list."
USE: "Please link the official source wherever stats appear."
USE: "This is catchy — keep it."
USE: "Either reduce word count or add relevant H3s."

NEVER: "Consider revising..." / "You might want to..." / "Great job overall, just a few tweaks..."

════════════════════════════════════════
LEAP SCHOLAR CONTEXT
════════════════════════════════════════

Audience: Indian students aged 18–30 planning to study abroad.
Destinations covered: USA, UK, Germany, Canada, Australia, Singapore, Ireland, Netherlands, France, Bulgaria, Finland, Switzerland.
Topics: Scholarships, Visa Guides, University Guides, Career Guides, Financial Planning, SOP Writing.
Student: ambitious, cost-conscious, first-generation or aspirational, needs trustworthy information for life-defining decisions.
Voice: mentor-like (not corporate), aspirational (not alarmist), specific (not generic), transparent (not promotional).
Goal: Every blog makes the student feel informed, empowered, and clear on their next step.

════════════════════════════════════════
EDGE CASES
════════════════════════════════════════

Blog <600 words → Flag: minimum 1,000 words required.
No FAQs → Flag: mandatory, suggest 4–5 PAA-style questions.
No meta title/description → Flag and write suggested versions in rewrite.
Multiple CTAs → Flag and remove all except conclusion CTA.
Key Takeaways present → Flag and remove entirely.
First-person phrases → Flag every instance and remove all.
"""


# ─────────────────────────────────────────────────────────────────────────────
# GOOGLE DOC FETCHER
# ─────────────────────────────────────────────────────────────────────────────
def extract_doc_id(url: str):
    for pattern in [r"/document/d/([a-zA-Z0-9_-]+)", r"id=([a-zA-Z0-9_-]+)"]:
        m = re.search(pattern, url)
        if m:
            return m.group(1)
    return None


def fetch_google_doc(url: str):
    doc_id = extract_doc_id(url)
    if not doc_id:
        raise ValueError(
            "Could not find a valid Google Doc ID in this URL.\n"
            "Make sure it's a standard Google Docs link (docs.google.com/document/d/...)."
        )

    # Export as HTML — preserves all hyperlinks unlike ?format=txt
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=html"
    try:
        resp = requests.get(export_url, timeout=30)
    except requests.exceptions.ConnectionError:
        raise ConnectionError("Network error — could not reach Google Docs. Check your internet connection.")

    if resp.status_code == 403:
        raise PermissionError(
            "Document is private.\n"
            "Please change sharing to 'Anyone with the link can view' and try again."
        )
    if resp.status_code != 200:
        raise ConnectionError(f"Could not fetch document (HTTP {resp.status_code}). Please check the link.")

    if not resp.text.strip():
        raise ValueError("The document appears to be empty.")

    # Parse HTML and inject [LINK: url] markers inline next to anchor text
    # This makes existing hyperlinks visible to the review engine as source citations
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup.find_all("a", href=True):
        href = tag.get("href", "").strip()
        link_text = tag.get_text()
        if not href or href.startswith("#") or href.startswith("javascript"):
            tag.replace_with(link_text)
            continue
        # Unwrap Google redirect URLs to get the real destination
        real_url = href
        if "google.com/url" in href:
            q_match = re.search(r'[?&]q=([^&]+)', href)
            if q_match:
                real_url = unquote(q_match.group(1))
        tag.replace_with(f"{link_text} [LINK: {real_url}]")

    text = soup.get_text(separator="\n")

    # Collapse excess blank lines (max 2 consecutive)
    lines = [l.strip() for l in text.splitlines()]
    cleaned_lines = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned_lines.append(line)
        else:
            blank_count = 0
            cleaned_lines.append(line)

    text = "\n".join(cleaned_lines).strip()

    if not text:
        raise ValueError("The document appears to be empty after processing.")

    if len(text) > 30000:
        raise ValueError(
            f"This document is very long ({len(text):,} characters). "
            "Please trim it to under ~30,000 characters to avoid token limit issues."
        )

    non_empty = [l for l in cleaned_lines if l.strip()]
    title = non_empty[0] if non_empty else "Untitled Blog"
    return text, title


# ─────────────────────────────────────────────────────────────────────────────
# OPENAI API
# ─────────────────────────────────────────────────────────────────────────────
def run_initial_review(api_key: str, blog_text: str, fact_check_text: str = "") -> str:
    client = OpenAI(api_key=api_key)
    today = datetime.now().strftime('%d %B %Y')

    # Build fact correction context for the rewrite if we have fact check results
    fact_correction_block = ""
    source_needed_block = ""
    if fact_check_text and fact_check_text.strip():
        # Extract INCORRECT and OUTDATED verdicts — these go as hard corrections
        corrections = []
        # Extract UNVERIFIABLE verdicts — these go as [SOURCE NEEDED: ...] markers
        source_needed = []
        lines = fact_check_text.split("\n")
        current_fact = current_verdict = current_detail = ""
        for line in lines:
            line = line.strip()
            if line.startswith("FACT:"):
                current_fact = line.replace("FACT:", "").strip()
            elif line.startswith("VERDICT:"):
                current_verdict = line.replace("VERDICT:", "").strip()
            elif line.startswith("DETAIL:"):
                current_detail = line.replace("DETAIL:", "").strip()
                if current_fact and current_verdict:
                    v = current_verdict.upper()
                    if "INCORRECT" in v or "OUTDATED" in v:
                        corrections.append(
                            f"• WRONG IN BLOG: \"{current_fact}\"\n"
                            f"  CORRECT FACT: {current_detail}"
                        )
                    elif "UNVERIFIABLE" in v:
                        source_needed.append(current_fact)
                current_fact = current_verdict = current_detail = ""

        if corrections:
            fact_correction_block = (
                "\n\nFACT CORRECTIONS — CRITICAL:\n"
                "The following facts in the blog have been verified as incorrect or outdated "
                "by an independent fact-checker. You MUST use the corrected figures in the "
                "rewritten blog. Do not use the original wrong figures under any circumstances.\n\n"
                + "\n\n".join(corrections)
            )

        if source_needed:
            source_needed_block = (
                "\n\nUNVERIFIABLE FACTS — ADD SOURCE MARKERS:\n"
                "The following claims could not be verified by the fact-checker. "
                "In the rewritten blog, add [SOURCE NEEDED: specific official source name] "
                "immediately after each of these claims. Name the specific source "
                "(e.g. 'Coventry University admissions page', 'UK Home Office', 'BAMF Germany') "
                "— never write a generic placeholder.\n\n"
                + "\n".join(f"• \"{f}\"" for f in source_needed)
            )

    prompt = (
        f"Today's date is {today}. Use this exact date as the REVIEW DATE in your output.\n\n"
        "Please review the following Leap Scholar blog in full.\n\n"
        "Apply the complete 5-step SOP (including Step 6 cross-consistency check) "
        "and all 10 pattern rules (K-1 through K-10).\n\n"
        "IMPORTANT: Anywhere you see [LINK: url] in the blog text, a hyperlink already "
        "exists there in the original Google Doc. Treat it as a valid source citation. "
        "Do NOT flag these as missing sources.\n\n"
        + fact_correction_block
        + source_needed_block
        + "\n\nProduce BOTH outputs separated by the EXACT markers:\n"
        "---REVIEW DOCUMENT START--- ... ---REVIEW DOCUMENT END---\n"
        "---REWRITTEN BLOG START--- ... ---REWRITTEN BLOG END---\n\n"
        "Here is the blog:\n\n"
        f"{blog_text}"
    )
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": MASTER_PROMPT},
            {"role": "user", "content": prompt},
        ],
        max_tokens=8000,
    )
    return response.choices[0].message.content


def run_followup(api_key: str, history: list, user_message: str) -> str:
    client = OpenAI(api_key=api_key)
    # history stored as OpenAI-format messages list
    messages = [{"role": "system", "content": MASTER_PROMPT}] + history + [
        {"role": "user", "content": user_message}
    ]
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        max_tokens=4000,
    )
    return response.choices[0].message.content


# ─────────────────────────────────────────────────────────────────────────────
# FACT CHECKER — Responses API with live web search
# Uses client.responses.create() NOT chat.completions — only endpoint that
# supports web_search_preview. Same API key, no extra signup needed.
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# FACT CHECKER — Serper.dev primary + GPT-4o Responses API fallback
# Architecture:
#   1. Extract individual facts from blog text using GPT-4o
#   2. For each fact: run targeted Serper search (your code controls the query)
#   3. Pass search results to GPT-4o for verdict only — cannot skip or use memory
#   4. On ANY Serper error (429, 403, 404, timeout, exhausted): fall back to
#      GPT-4o Responses API with web_search_preview
#   5. Always tell user which mode ran via fact_check_mode return value
# ─────────────────────────────────────────────────────────────────────────────

SERPER_ENDPOINT = "https://google.serper.dev/search"
SERPER_RATE_LIMIT_DELAY = 1.2   # seconds between calls — stays within free tier rate limits
SERPER_MAX_RETRIES = 2          # retry once on 429 before rotating to fallback


FACT_EXTRACTION_PROMPT = """
You are a fact extraction assistant. Read the blog text below and extract every specific,
verifiable factual claim — numbers, statistics, percentages, fees, deadlines, salary thresholds,
rankings, policy rules, dates, and named figures.

Return ONLY a JSON array of strings. Each string = one fact exactly as stated in the blog.
Maximum 25 facts. Skip opinions, general statements, narrative, and CTAs.
Only extract claims that can be verified against an external source.

Example output:
["India is the largest source country for international students in Germany",
 "The blocked account requirement for 2026 is €11,904",
 "EU Blue Card minimum salary for STEM is €45,000+",
 "International students can work up to 140 full days per year"]

Return ONLY the JSON array. No preamble, no explanation, no markdown code fences.
"""

FACT_VERDICT_PROMPT = """
You are a professional fact-checker. You will be given:
1. A specific factual claim from a blog
2. Search results from Google (via Serper) relevant to that claim

Your job: read the search results carefully and give a verdict.

TODAY IS MARCH 2026. All verdicts must reflect what is TRUE as of early 2026.

CRITICAL RULES:
- Base your verdict ONLY on the search results provided — do not use training memory
- If multiple results conflict, use the most recently dated source
- For German policy facts (salaries, fees, thresholds): prefer official sources
  (make-it-in-germany.com, daad.de, bamf.de, study-in-germany.de)
- IMPORTANT: When a claim mentions a specific category (e.g. STEM, shortage occupation),
  verify the figure for THAT category specifically, not the general threshold
- Germany updates salary thresholds every January — a 2023 figure is NOT valid for 2026

VERDICT OPTIONS:
✅ VERIFIED — search results confirm the claim is accurate for 2026
⚠️ PARTIALLY CORRECT — figure is close but slightly off
⚠️ OUTDATED — was correct before but search results show it has changed
🔴 INCORRECT — search results clearly contradict the claim
⚠️ UNVERIFIABLE — search results do not contain enough information to verify

OUTPUT FORMAT (use exactly this, nothing else):
VERDICT: [one of the 5 options above]
DETAIL: [1-2 sentences. What the search results show. If incorrect, state the correct 2026 figure.]
SOURCE: [URL of the most relevant search result used]
"""

FACT_CHECK_FALLBACK_PROMPT = """
You are a professional fact-checker specialising in education, study abroad, immigration policy, and international universities.

TODAY'S DATE IS MARCH 2026. All verdicts must reflect what is TRUE as of early 2026.

Search the web to verify every specific claim, number, statistic, fee, deadline, ranking, or policy in the blog.

OUTPUT FORMAT — use EXACTLY these markers:

---FACT CHECK START---

FACT CHECK SUMMARY:
Total facts checked: [number]
✅ Verified: [number]
⚠️ Partially correct / outdated: [number]
🔴 Incorrect: [number]
⚠️ Unverifiable: [number]

---FACT CHECK ITEMS---

FACT: [Exact quote from blog]
VERDICT: [✅ VERIFIED / ⚠️ PARTIALLY CORRECT / ⚠️ OUTDATED / 🔴 INCORRECT / ⚠️ UNVERIFIABLE]
DETAIL: [What your live search found. Always state the current 2026 figure. Max 2 sentences.]
SOURCE: [URL from your search result]

[Repeat for every fact]

---FACT CHECK END---

CRITICAL SEARCH RULES:
- For every salary, fee, threshold, policy — search with "2026" in the query
- For German facts: prefer make-it-in-germany.com, daad.de, bamf.de, study-in-germany.de
- When a claim mentions STEM or shortage occupations, verify the STEM/shortage threshold specifically
- Trust the most recently published source. Discard sources older than 12 months for policy facts.
- NEVER flag a stat that has [LINK: url] next to it — it is already sourced.
"""


def extract_facts_from_blog(api_key: str, blog_text: str) -> list:
    """Use GPT-4o to extract a clean list of verifiable facts from the blog."""
    client = OpenAI(api_key=api_key)
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": FACT_EXTRACTION_PROMPT},
                {"role": "user", "content": f"Extract all verifiable facts from this blog:\n\n{blog_text}"}
            ],
            max_tokens=1500,
            temperature=0,
        )
        raw = response.choices[0].message.content.strip()
        # Strip markdown fences if model wraps in ```json
        raw = re.sub(r'^```[a-z]*\n?', '', raw, flags=re.MULTILINE)
        raw = raw.replace('```', '').strip()
        facts = __import__('json').loads(raw)
        if isinstance(facts, list):
            return [str(f) for f in facts[:25]]
    except Exception:
        pass
    return []


def serper_search(query: str, serper_key: str, retries: int = 0) -> dict | None:
    """
    Run one Serper search. Returns the parsed JSON response or None on failure.
    Handles rate limiting with exponential backoff.
    Raises SerperExhaustedError on quota errors (403) so caller can fall back.
    """
    headers = {
        "X-API-KEY": serper_key,
        "Content-Type": "application/json",
    }
    payload = {"q": query, "num": 5}
    try:
        resp = requests.post(
            SERPER_ENDPOINT, headers=headers,
            json=payload, timeout=10
        )
        if resp.status_code == 200:
            return resp.json()
        elif resp.status_code == 429:
            # Rate limited — wait and retry once
            if retries < SERPER_MAX_RETRIES:
                time.sleep(3 * (retries + 1))
                return serper_search(query, serper_key, retries + 1)
            else:
                raise SerperExhaustedError("Serper rate limit hit after retries.")
        elif resp.status_code in (403, 401):
            raise SerperExhaustedError(f"Serper quota exhausted or key invalid (HTTP {resp.status_code}).")
        else:
            # 404, 500, etc — treat as transient, return None so this fact is skipped
            return None
    except SerperExhaustedError:
        raise
    except Exception:
        return None


class SerperExhaustedError(Exception):
    """Raised when Serper quota is exhausted or key is invalid — triggers fallback."""
    pass


def build_serper_query(fact: str) -> str:
    """
    Build a targeted, year-anchored search query for a given fact.
    Always includes 2026 for policy/salary facts to avoid stale results.
    """
    # Keywords that need year-anchored queries
    year_anchored_keywords = [
        "salary", "threshold", "blocked account", "semester", "fee", "tuition",
        "stipend", "scholarship", "blue card", "visa", "permit", "working days",
        "source country", "ranked", "ranking", "students in germany", "shortage",
        "minimum", "required amount", "budget", "cost", "exchange rate",
    ]
    needs_year = any(kw in fact.lower() for kw in year_anchored_keywords)
    query = fact.strip()
    if needs_year and "2026" not in query and "2025" not in query:
        query = f"{query} 2026"
    return query


def verdict_from_search_results(
    api_key: str, fact: str, search_results: dict
) -> tuple[str, str, str]:
    """
    Ask GPT-4o to give a verdict based ONLY on the provided search results.
    Returns (verdict_line, detail_line, source_line).
    GPT-4o cannot use memory here — it only reads what we pass it.
    """
    # Format top 5 results as readable text
    snippets = []
    organic = search_results.get("organic", [])
    for i, r in enumerate(organic[:5]):
        title = r.get("title", "")
        snippet = r.get("snippet", "")
        link = r.get("link", "")
        date = r.get("date", "")
        date_str = f" [{date}]" if date else ""
        snippets.append(f"Result {i+1}{date_str}: {title}\n{snippet}\nURL: {link}")

    if not snippets:
        return (
            "⚠️ UNVERIFIABLE",
            "No search results returned for this claim.",
            "N/A"
        )

    search_text = "\n\n".join(snippets)

    client = OpenAI(api_key=api_key)
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": FACT_VERDICT_PROMPT},
                {"role": "user", "content": (
                    f"CLAIM TO VERIFY: {fact}\n\n"
                    f"SEARCH RESULTS:\n{search_text}\n\n"
                    "Give your verdict using the exact format specified."
                )}
            ],
            max_tokens=300,
            temperature=0,
        )
        raw = response.choices[0].message.content.strip()
        verdict = "⚠️ UNVERIFIABLE"
        detail = "Could not parse verdict."
        source = "N/A"
        for line in raw.split("\n"):
            line = line.strip()
            if line.startswith("VERDICT:"):
                verdict = line.replace("VERDICT:", "").strip()
            elif line.startswith("DETAIL:"):
                detail = line.replace("DETAIL:", "").strip()
            elif line.startswith("SOURCE:"):
                source = line.replace("SOURCE:", "").strip()
        return verdict, detail, source
    except Exception as e:
        return "⚠️ UNVERIFIABLE", f"Verdict API error: {e}", "N/A"


def run_fact_check_serper(
    api_key: str, serper_key: str, blog_text: str
) -> tuple[str, str]:
    """
    PRIMARY FACT CHECKER — Serper-powered.
    Returns (fact_check_text, mode) where mode is 'high_accuracy' or 'standard'.

    Flow:
    1. Extract facts from blog using GPT-4o
    2. For each fact: build targeted query → Serper search → GPT-4o verdict
    3. Rate limit: 1.2s between Serper calls
    4. On SerperExhaustedError: immediately fall back to GPT-4o Responses API
    5. On any other error per-fact: mark as UNVERIFIABLE and continue
    """
    # Step 1: extract facts
    facts = extract_facts_from_blog(api_key, blog_text)
    if not facts:
        # Can't extract facts — fall back entirely
        return run_fact_check_gpt_fallback(api_key, blog_text), "standard"

    results = []
    verified = partially = outdated = incorrect = unverifiable = 0

    for i, fact in enumerate(facts):
        try:
            query = build_serper_query(fact)
            search_data = serper_search(query, serper_key)

            if search_data is None:
                # Transient error — mark unverifiable and continue
                results.append({
                    "fact": fact,
                    "verdict": "⚠️ UNVERIFIABLE",
                    "detail": "Search returned no results for this claim.",
                    "source": "N/A"
                })
                unverifiable += 1
            else:
                verdict, detail, source = verdict_from_search_results(
                    api_key, fact, search_data
                )
                results.append({
                    "fact": fact,
                    "verdict": verdict,
                    "detail": detail,
                    "source": source
                })
                v = verdict.upper()
                if "VERIFIED" in v: verified += 1
                elif "PARTIALLY" in v: partially += 1
                elif "OUTDATED" in v: outdated += 1
                elif "INCORRECT" in v: incorrect += 1
                else: unverifiable += 1

            # Rate limit — stay within Serper's limits
            if i < len(facts) - 1:
                time.sleep(SERPER_RATE_LIMIT_DELAY)

        except SerperExhaustedError:
            # Quota exhausted — fall back to GPT-4o for the entire check
            fallback_text = run_fact_check_gpt_fallback(api_key, blog_text)
            return fallback_text, "standard"

    # Build structured output matching the existing parser format
    total = len(results)
    lines = [
        "",
        "FACT CHECK SUMMARY:",
        f"Total facts checked: {total}",
        f"✅ Verified: {verified}",
        f"⚠️ Partially correct / outdated: {partially + outdated}",
        f"🔴 Incorrect: {incorrect}",
        f"⚠️ Unverifiable: {unverifiable}",
        "",
        "---FACT CHECK ITEMS---",
        "",
    ]
    for r in results:
        lines += [
            f"FACT: {r['fact']}",
            f"VERDICT: {r['verdict']}",
            f"DETAIL: {r['detail']}",
            f"SOURCE: {r['source']}",
            "",
        ]
    return "\n".join(lines), "high_accuracy"


def run_fact_check_gpt_fallback(api_key: str, blog_text: str) -> str:
    """
    FALLBACK FACT CHECKER — GPT-4o Responses API with web_search_preview.
    Used when Serper is exhausted or unavailable.
    Returns raw fact check text (to be parsed by parse_fact_check).
    """
    client = OpenAI(api_key=api_key)
    response = client.responses.create(
        model="gpt-4o",
        tools=[{"type": "web_search_preview"}],
        input=[
            {"role": "system", "content": FACT_CHECK_FALLBACK_PROMPT},
            {"role": "user", "content": (
                "Today is March 2026. Fact-check every specific number, statistic, "
                "percentage, fee, deadline, ranking, or policy claim in this blog.\n\n"
                "For every search, include '2026' in your query to get current results. "
                "For STEM/shortage occupation salary claims, verify the STEM threshold "
                "specifically, not the general threshold.\n\n"
                "Trust the most recently published official source.\n\n"
                f"Blog text:\n\n{blog_text}"
            )}
        ],
    )
    raw = response.output_text
    # If it returns with markers, extract just the inner content
    m = re.search(r"---FACT CHECK START---(.*?)---FACT CHECK END---", raw, re.DOTALL)
    if m:
        return m.group(1).strip()
    return raw.strip()


def run_fact_check(
    api_key: str, blog_text: str, serper_key: str = ""
) -> tuple[str, str]:
    """
    Master fact check dispatcher.
    Returns (fact_check_text, mode).
    mode = 'high_accuracy' (Serper) or 'standard' (GPT-4o fallback).
    """
    if serper_key and serper_key.strip():
        try:
            return run_fact_check_serper(api_key, serper_key.strip(), blog_text)
        except SerperExhaustedError:
            pass
        except Exception:
            pass
    # No Serper key or Serper completely failed
    return run_fact_check_gpt_fallback(api_key, blog_text), "standard"


def parse_fact_check(text: str) -> str:
    """Extract content between fact check markers if present, else return as-is."""
    m = re.search(r"---FACT CHECK START---(.*?)---FACT CHECK END---", text, re.DOTALL)
    if m:
        return m.group(1).strip()
    return text.strip()


# ─────────────────────────────────────────────────────────────────────────────
# PARSE OPENAI RESPONSE
# FIX: Instead of a broken character-split fallback, surface a clear error
# ─────────────────────────────────────────────────────────────────────────────
def parse_response(text: str):
    review, rewrite = "", ""
    m1 = re.search(r"---REVIEW DOCUMENT START---(.*?)---REVIEW DOCUMENT END---", text, re.DOTALL)
    if m1:
        review = m1.group(1).strip()
    m2 = re.search(r"---REWRITTEN BLOG START---(.*?)---REWRITTEN BLOG END---", text, re.DOTALL)
    if m2:
        rewrite = m2.group(1).strip()

    if not review and not rewrite:
        raise ValueError(
            "The model did not return the expected output markers.\n\n"
            "This can happen if the blog is too complex or the model hit a limit. "
            "Please try again or use '🔄 New Review' to reset."
        )
    if not review:
        review = "⚠️ Review section was not returned by the model. Please try again."
    if not rewrite:
        rewrite = "⚠️ Rewritten blog was not returned by the model. Please try again."

    return review, rewrite


# ─────────────────────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def add_hr(doc, color="1a56a0"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)


def mixed_run(para, text):
    """Render **bold** and *italic* markdown inline."""
    # Handle bold first, then italic
    segments = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in segments:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = para.add_run(part[1:-1])
            r.italic = True
        else:
            r = para.add_run(part)
        r.font.size = Pt(11)


# ─────────────────────────────────────────────────────────────────────────────
# SAFE FILENAME HELPER
# FIX: handles unicode/international characters in blog titles
# ─────────────────────────────────────────────────────────────────────────────
def safe_filename(title: str, max_len: int = 40) -> str:
    normalized = unicodedata.normalize('NFKD', title)
    ascii_title = normalized.encode('ascii', 'ignore').decode('ascii')
    return re.sub(r'[^\w\s-]', '', ascii_title)[:max_len].strip()


# ─────────────────────────────────────────────────────────────────────────────
# BUILD REVIEW DOCX
# ─────────────────────────────────────────────────────────────────────────────
def build_review_docx(review_text: str, blog_title: str, fact_check_text: str = "", fact_check_mode: str = "standard") -> bytes:
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.2)
        sec.right_margin = Inches(1.2)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LEAP SCHOLAR — BLOG REVIEW")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Editorial Review by Krutika AI")
    r2.italic = True; r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    add_hr(doc, "1a56a0")
    doc.add_paragraph()

    in_scorecard = False
    in_priority = False

    for raw_line in review_text.split('\n'):
        s = raw_line.strip()
        if not s:
            doc.add_paragraph(); continue

        if s == "---SCORECARD---":
            in_scorecard = True
            h = doc.add_heading("📊 Scorecard", 1)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
            continue
        if s == "---MACRO SUMMARY---":
            in_scorecard = False
            h = doc.add_heading("📝 Macro Summary", 1)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
            continue
        if s == "---SECTION-WISE REVIEW---":
            h = doc.add_heading("🔍 Section-Wise Review", 1)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
            continue
        if s == "---SEO AUDIT---":
            h = doc.add_heading("🔎 SEO Audit", 1)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
            continue
        if s == "---GRAMMAR & STYLE AUDIT---":
            h = doc.add_heading("✏️ Grammar & Style Audit", 1)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
            continue
        if s == "---PRIORITY ACTION LIST---":
            in_priority = True
            h = doc.add_heading("⚡ Priority Action List", 1)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
            continue

        if s.startswith("BLOG TITLE:"):
            p = doc.add_paragraph()
            r1 = p.add_run("Blog Title:  "); r1.bold = True; r1.font.size = Pt(12)
            p.add_run(s.replace("BLOG TITLE:", "").strip()).font.size = Pt(12)

        elif s.startswith("REVIEW DATE:") or s.startswith("REVIEWER:"):
            p = doc.add_paragraph()
            r = p.add_run(s); r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        elif s.startswith("OVERALL STATUS:"):
            val = s.replace("OVERALL STATUS:", "").strip()
            p = doc.add_paragraph()
            r1 = p.add_run("Overall Status:  "); r1.bold = True; r1.font.size = Pt(12)
            r2 = p.add_run(val); r2.bold = True; r2.font.size = Pt(12)
            if "APPROVED" in val and "MINOR" not in val:
                r2.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
            elif "MINOR" in val:
                r2.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
            else:
                r2.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

        elif s.startswith("OVERALL SCORE:"):
            p = doc.add_paragraph()
            r = p.add_run(s); r.bold = True; r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
            p.paragraph_format.space_before = Pt(6)

        elif s.startswith("SECTION:"):
            doc.add_paragraph()
            h = doc.add_heading(s.replace("SECTION:", "Section:"), 2)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)

        elif s.startswith("STATUS:"):
            val = s.replace("STATUS:", "").strip()
            p = doc.add_paragraph()
            r = p.add_run(f"Status: {val}"); r.bold = True; r.font.size = Pt(11)
            if "🔴" in val: r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
            elif "⚠️" in val: r.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
            elif "✅" in val: r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)

        elif s.startswith("ISSUES FOUND:"):
            p = doc.add_paragraph()
            r = p.add_run("Issues Found:"); r.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x2e, 0x40, 0x57)

        elif s.startswith("🔴"):
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(s); r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

        elif s.startswith("⚠️"):
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(s); r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)

        elif s.startswith("✅"):
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(s); r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)

        elif s.startswith("Fix:"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_before = Pt(2)
            r1 = p.add_run("Fix: "); r1.bold = True; r1.font.size = Pt(11)
            r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)
            p.add_run(s.replace("Fix:", "").strip()).font.size = Pt(11)

        elif s.startswith("Original:"):
            p = doc.add_paragraph()
            r1 = p.add_run("Original: "); r1.bold = True; r1.font.size = Pt(11)
            r2 = p.add_run(s.replace("Original:", "").strip())
            r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

        elif in_priority and re.match(r'^\d+\.', s):
            p = doc.add_paragraph(style='List Number')
            r = p.add_run(s); r.font.size = Pt(11)
            if "[CRITICAL]" in s: r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
            elif "[HIGH]" in s: r.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)

        elif in_scorecard and re.match(r'^\d+\.', s):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            parts = s.split(":", 1)
            if len(parts) == 2:
                r1 = p.add_run(parts[0] + ":"); r1.bold = True; r1.font.size = Pt(11)
                p.add_run(parts[1]).font.size = Pt(11)
            else:
                p.add_run(s).font.size = Pt(11)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.add_run(s).font.size = Pt(11)

    # ── 🔬 Fact Check Report — appended at end in purple ──
    if fact_check_text and fact_check_text.strip():
        doc.add_paragraph()
        add_hr(doc, "7c3aed")
        doc.add_paragraph()
        h = doc.add_heading("🔬 Fact Check Report", 1)
        if h.runs: h.runs[0].font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)

        # Mode badge
        mode_p = doc.add_paragraph()
        if fact_check_mode == "high_accuracy":
            mode_r = mode_p.add_run("🟢 High Accuracy Mode — Serper (Google) powered search, one targeted search per fact")
            mode_r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
        else:
            mode_r = mode_p.add_run("🟡 Standard Mode — GPT-4o web search (add Serper API key for higher accuracy)")
            mode_r.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
        mode_r.font.size = Pt(10); mode_r.bold = True

        sub = doc.add_paragraph()
        sub_r = sub.add_run("Independent verification of all facts, figures, and statistics against live web sources.")
        sub_r.italic = True; sub_r.font.size = Pt(10)
        sub_r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
        doc.add_paragraph()
        for raw_line in fact_check_text.split('\n'):
            s = raw_line.strip()
            if not s: doc.add_paragraph(); continue
            if s == "---FACT CHECK ITEMS---": continue
            elif s.startswith("FACT CHECK SUMMARY:"):
                p = doc.add_paragraph()
                r = p.add_run("Summary"); r.bold = True; r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)
            elif any(s.startswith(x) for x in ["Total facts", "✅ Verified", "⚠️ Partially", "🔴 Incorrect", "⚠️ Unverifiable"]):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run(s); r.font.size = Pt(11)
                if s.startswith("✅"): r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
                elif s.startswith("⚠️"): r.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
                elif s.startswith("🔴"): r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
            elif s.startswith("FACT:"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                r1 = p.add_run("FACT: "); r1.bold = True; r1.font.size = Pt(11)
                r1.font.color.rgb = RGBColor(0x2e, 0x40, 0x57)
                p.add_run(s.replace("FACT:", "").strip()).font.size = Pt(11)
            elif s.startswith("VERDICT:"):
                val = s.replace("VERDICT:", "").strip()
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run(f"Verdict: {val}"); r.bold = True; r.font.size = Pt(11)
                if "VERIFIED" in val: r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
                elif "INCORRECT" in val: r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
                else: r.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
            elif s.startswith("DETAIL:"):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
                r1 = p.add_run("Detail: "); r1.bold = True; r1.font.size = Pt(10)
                r1.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                p.add_run(s.replace("DETAIL:", "").strip()).font.size = Pt(10)
            elif s.startswith("SOURCE:"):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
                r1 = p.add_run("Source: "); r1.bold = True; r1.font.size = Pt(10)
                r1.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                p.add_run(s.replace("SOURCE:", "").strip()).font.size = Pt(10)
            else:
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
                p.add_run(s).font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# BUILD REWRITTEN BLOG DOCX
# ─────────────────────────────────────────────────────────────────────────────
def build_rewritten_docx(rewritten_text: str, blog_title: str) -> bytes:
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.2)
        sec.right_margin = Inches(1.2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REWRITTEN BLOG"); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Revised by Krutika AI — Leap Scholar Editorial System")
    r2.italic = True; r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    add_hr(doc, "22c55e")
    doc.add_paragraph()

    for raw_line in rewritten_text.split('\n'):
        s = raw_line.strip()
        if not s:
            doc.add_paragraph(); continue

        if s.startswith("# "):
            h = doc.add_heading(s[2:].strip(), 1)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)

        elif s.startswith("## "):
            h = doc.add_heading(s[3:].strip(), 2)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)

        elif s.startswith("### "):
            h = doc.add_heading(s[4:].strip(), 3)
            if h.runs: h.runs[0].font.color.rgb = RGBColor(0x2e, 0x40, 0x57)

        elif s.startswith("- ") or s.startswith("• "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            mixed_run(p, s[2:].strip())

        elif re.match(r'^\d+\.\s', s):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            mixed_run(p, re.sub(r'^\d+\.\s*', '', s))

        elif s.startswith("[SOURCE NEEDED"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(s); r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

        elif s.startswith("[TABLE RECOMMENDED"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(s); r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x1a, 0x56, 0xa0)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            mixed_run(p, s)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────
defaults = {
    "messages": [], "phase": "home", "doc_url": "",
    "blog_text": "", "blog_title": "",
    "review_bytes": None, "rewrite_bytes": None,
    "openai_history": [], "review_done": False,
    "followup_count": 0,
    "fact_check_mode": "high_accuracy",  # 'high_accuracy' | 'standard'
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

MAX_FOLLOWUPS = 10


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# FIX: st.secrets fallback so key can be pre-set on Streamlit Cloud
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="font-family:'Syne',sans-serif;font-weight:700;font-size:1.1rem;
                color:#22c55e;margin-bottom:16px;margin-top:8px;">✦ Configuration</div>
    """, unsafe_allow_html=True)

    default_key = st.secrets.get("OPENAI_API_KEY", "") if hasattr(st, "secrets") else ""
    api_key = st.text_input(
        "OpenAI API Key", type="password", placeholder="sk-...",
        value=default_key,
        help="Get your key at platform.openai.com/api-keys"
    )

    default_serper = st.secrets.get("SERPER_API_KEY", "") if hasattr(st, "secrets") else ""
    serper_key = st.text_input(
        "Serper API Key", type="password", placeholder="serper key...",
        value=default_serper,
        help="Get a free key at serper.dev — 2,500 free searches/month (~125 blogs)"
    )

    # Mode indicator
    if serper_key and serper_key.strip():
        st.markdown("""
        <div style="font-size:0.75rem;color:#22c55e;margin-top:6px;
                    background:rgba(34,197,94,0.08);border:1px solid rgba(34,197,94,0.2);
                    border-radius:8px;padding:6px 10px;">
            🟢 <strong>High Accuracy Mode</strong><br>
            Fact checker uses Serper (Google) search
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="font-size:0.75rem;color:#d97706;margin-top:6px;
                    background:rgba(217,119,6,0.08);border:1px solid rgba(217,119,6,0.2);
                    border-radius:8px;padding:6px 10px;">
            🟡 <strong>Standard Mode</strong><br>
            Add Serper key for higher accuracy
        </div>
        """, unsafe_allow_html=True)

    st.markdown("""
    <div style="font-size:0.75rem;color:#6b7280;margin-top:8px;line-height:1.7;">
        Keys are never stored or logged.<br>
        Uses <strong style="color:#22c55e">gpt-4o</strong><br><br>
        📄 Google Doc must be set to<br><em>"Anyone with the link can view"</em>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    if st.button("🔄 New Review", use_container_width=True):
        for k, v in defaults.items():
            st.session_state[k] = v
        st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# HOME SCREEN
# ─────────────────────────────────────────────────────────────────────────────
if st.session_state.phase == "home":

    st.markdown("""
    <div class="hero-wrap">
        <div class="hero-logo">
            <img src="https://d14lg9nzq1d3lc.cloudfront.net/advance-website/assets/images/company-logo/logo.svg"
                 style="width:200px;height:90px;object-fit:contain;filter:brightness(0) invert(1);" alt="Leap Finance" />
        </div>
        <div class="hero-title">Ready for your<br><span>Blog Review?</span></div>
        <p class="hero-sub" style="text-align:center;width:100%;display:block;margin-left:auto;margin-right:auto;">Paste a Google Doc link and I'll review it in Krutika's style,<br>
        section by section, with fixes and a fully rewritten version.</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="status-badge">
        <span class="status-dot"></span>&nbsp; Krutika AI · Ready
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="cards-row">
        <div class="card">
            <div class="card-icon">📋</div>
            <div class="card-title">Section-Wise Review</div>
            <div class="card-desc">Every H2 and H3 reviewed individually with 🔴 issues and ✅ fixes, exactly how Krutika does it.</div>
        </div>
        <div class="card">
            <div class="card-icon">✍️</div>
            <div class="card-title">Rewritten Blog</div>
            <div class="card-desc">A fully rewritten version of your blog incorporating every fix, download-ready as a Word doc.</div>
        </div>
        <div class="card">
            <div class="card-icon">📊</div>
            <div class="card-title">Scorecard + SEO Audit</div>
            <div class="card-desc">10-category scorecard, SEO compliance check, grammar audit, and a priority action list.</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="home-input-wrap">', unsafe_allow_html=True)
    col1, col2 = st.columns([5, 1])
    with col1:
        url = st.text_input(
            "", placeholder="✦  Paste your Google Doc link here...",
            key="url_input", label_visibility="collapsed"
        )
    with col2:
        go = st.button("Review →", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

    if go:
        if not api_key:
            st.error("⚠️ Please enter your OpenAI API key in the sidebar first.")
        elif not url:
            st.markdown("""
            <div style="text-align:center;color:#ef4444;font-size:0.85rem;margin-top:12px;">
                Please paste a Google Doc link before clicking Review →
            </div>""", unsafe_allow_html=True)
        else:
            with st.spinner("Fetching your Google Doc..."):
                try:
                    blog_text, blog_title = fetch_google_doc(url)
                    st.session_state.update({
                        "doc_url": url, "blog_text": blog_text,
                        "blog_title": blog_title, "phase": "chat",
                        "review_done": False, "followup_count": 0,
                        "fact_check_mode": "high_accuracy",
                        "messages": [{
                            "role": "ai",
                            "content": (
                                f"✦ Document fetched: **{blog_title}**\n\n"
                                "Running fact check + full editorial review...\n\n"
                                "I'll produce:\n"
                                "**① Section-wise Review Document** with 🔬 Fact Check Report (.docx)\n"
                                "**② Rewritten Blog** with corrected facts incorporated (.docx)\n\n"
                                "This will take 60–120 seconds. Hang tight."
                            )
                        }]
                    })
                    st.rerun()
                except (ValueError, PermissionError, ConnectionError) as e:
                    st.error(f"❌ {e}")
                except Exception as e:
                    st.error(f"❌ Unexpected error: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# CHAT SCREEN
# ─────────────────────────────────────────────────────────────────────────────
else:

    st.markdown("""
    <div style="padding: 28px 0 8px; display:flex; align-items:center; gap:12px;">
        <div style="width:32px;height:32px;background:linear-gradient(135deg,#22c55e,#16a34a);
                    border-radius:10px;display:flex;align-items:center;justify-content:center;
                    overflow:hidden;box-shadow:0 0 12px rgba(34,197,94,0.3);">
            <img src="https://d14lg9nzq1d3lc.cloudfront.net/advance-website/assets/images/company-logo/logo.svg"
                 style="width:22px;height:22px;object-fit:contain;filter:brightness(0) invert(1);" alt="Leap" />
        </div>
        <div>
            <div style="font-family:'Syne',sans-serif;font-weight:700;font-size:1rem;color:#f0f0f0;">
                Krutika AI</div>
            <div style="font-size:0.72rem;color:#22c55e;letter-spacing:0.05em;text-transform:uppercase;font-weight:500;">
                <span style="display:inline-block;width:5px;height:5px;background:#22c55e;border-radius:50%;
                             margin-right:5px;vertical-align:middle;animation:pulse-dot 2s infinite;"></span>
                Reviewing · Active</div>
        </div>
    </div>
    <div class="section-divider"></div>
    """, unsafe_allow_html=True)

    if st.session_state.doc_url:
        short = (st.session_state.doc_url[:55] + "...") if len(st.session_state.doc_url) > 55 else st.session_state.doc_url
        st.markdown(f"""
        <div style="display:flex;align-items:center;gap:8px;margin-bottom:20px;
                    background:rgba(255,255,255,0.03);border:1px solid rgba(255,255,255,0.07);
                    border-radius:10px;padding:10px 14px;width:fit-content;max-width:100%;">
            <span style="font-size:13px;">🔗</span>
            <span style="font-size:0.78rem;color:#6b7280;font-family:'DM Sans',sans-serif;
                         overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">{short}</span>
        </div>
        """, unsafe_allow_html=True)

    # ── AUTO-TRIGGER: 3-step flow ──
    # Step 1: Fact check (Serper or GPT-4o fallback)
    # Step 2: Review + rewrite WITH fact corrections injected
    # Step 3: Build docx files
    if not st.session_state.review_done and st.session_state.blog_text:
        with st.spinner("🔬 Step 1/3 — Fact-checking with live search..."):
            try:
                fact_check_text, fc_mode = run_fact_check(
                    api_key, st.session_state.blog_text,
                    serper_key if 'serper_key' in dir() else ""
                )
                st.session_state.fact_check_mode = fc_mode
            except Exception as e:
                fact_check_text = f"Fact check unavailable: {e}"
                st.session_state.fact_check_mode = "standard"

        with st.spinner("✍️ Step 2/3 — Running editorial review + incorporating corrections..."):
            try:
                # Pass fact check results so corrections are baked into the rewrite
                raw = run_initial_review(
                    api_key,
                    st.session_state.blog_text,
                    fact_check_text
                )
                review_text, rewritten_text = parse_response(raw)
            except Exception as e:
                st.session_state.review_done = True
                st.session_state.messages.append({
                    "role": "ai",
                    "content": f"❌ Error during review: {e}\n\nCheck your API key in the sidebar and use '🔄 New Review' to reset."
                })
                st.rerun()

        with st.spinner("📄 Step 3/3 — Building documents..."):
            try:
                st.session_state.review_bytes = build_review_docx(
                    review_text, st.session_state.blog_title, fact_check_text,
                    st.session_state.fact_check_mode
                )
                st.session_state.rewrite_bytes = build_rewritten_docx(
                    rewritten_text, st.session_state.blog_title
                )
                st.session_state.review_done = True
                st.session_state.openai_history = [
                    {"role": "user", "content": f"Review this Leap Scholar blog:\n\n{st.session_state.blog_text}"},
                    {"role": "assistant", "content": raw},
                ]

                # Build completion message with mode indicator
                mode = st.session_state.fact_check_mode
                if mode == "high_accuracy":
                    mode_badge = "🟢 **High Accuracy Mode** — Fact check powered by Serper (Google search)"
                else:
                    mode_badge = "🟡 **Standard Mode** — Fact check powered by GPT-4o web search\n\n⚠️ Add a Serper API key in the sidebar for higher accuracy on 2025/2026 figures."

                st.session_state.messages.append({
                    "role": "ai",
                    "content": (
                        f"✅ Review + Fact Check complete for **{st.session_state.blog_title}**\n\n"
                        f"{mode_badge}\n\n"
                        "Both documents are ready — download them below.\n\n"
                        "**What's been done:**\n"
                        "- 🔬 Every stat fact-checked against live search results\n"
                        "- ✍️ Corrected facts incorporated into the rewritten blog\n"
                        "- 📋 Section-wise review with fixes in the Review Document\n\n"
                        "Ask follow-up questions:\n"
                        "- *'Rewrite only the introduction'*\n"
                        "- *'Give me 5 better FAQ questions'*\n"
                        "- *'Make the conclusion more punchy'*"
                    )
                })
                st.rerun()
            except Exception as e:
                st.session_state.review_done = True
                st.session_state.messages.append({
                    "role": "ai",
                    "content": f"❌ Error building documents: {e}\n\nUse '🔄 New Review' to reset."
                })
                st.rerun()

    # ── Chat messages ──
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(f"""
            <div class="chat-msg-user">
                <div class="bubble">{msg["content"]}</div>
            </div>""", unsafe_allow_html=True)
        else:
            html = re.sub(r'\*\*(.*?)\*\*', r'<strong style="color:#d1fae5;">\1</strong>', msg["content"])
            html = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html)
            html = html.replace("\n", "<br>")
            st.markdown(f"""
            <div class="chat-msg-ai">
                <div class="ai-avatar" style="display:flex;align-items:center;justify-content:center;overflow:hidden;">
                    <img src="https://d14lg9nzq1d3lc.cloudfront.net/advance-website/assets/images/company-logo/logo.svg"
                         style="width:18px;height:18px;object-fit:contain;filter:brightness(0) invert(1);" alt="Leap" />
                </div>
                <div class="bubble">{html}</div>
            </div>""", unsafe_allow_html=True)

    # ── Download panel ──
    if st.session_state.review_done and st.session_state.review_bytes:
        st.markdown("""
        <div style="background:rgba(34,197,94,0.05);border:1px solid rgba(34,197,94,0.15);
                    border-radius:16px;padding:20px 24px;margin:8px 0 16px;">
            <div style="font-family:'Syne',sans-serif;font-weight:600;font-size:0.9rem;
                        color:#22c55e;margin-bottom:14px;letter-spacing:0.02em;">
                📥 Your Documents Are Ready
            </div>
        """, unsafe_allow_html=True)

        safe = safe_filename(st.session_state.blog_title)
        d1, d2 = st.columns(2)
        with d1:
            st.download_button(
                "📋 Download Review + Fact Check (.docx)",
                data=st.session_state.review_bytes,
                file_name=f"Review_{safe}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, key="dl_review"
            )
        with d2:
            st.download_button(
                "✍️ Download Rewritten Blog (.docx)",
                data=st.session_state.rewrite_bytes,
                file_name=f"Rewritten_{safe}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, key="dl_rewrite"
            )
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='height:120px'></div>", unsafe_allow_html=True)

    # ── Follow-up input (single functional input — decorative duplicate removed) ──
    # FIX: guard added so follow-up only fires after review is complete
    with st.container():
        c1, c2 = st.columns([6, 1])
        with c1:
            follow_up = st.text_input(
                "", placeholder="Ask a follow-up — e.g. 'Rewrite only the intro' or 'Make the conclusion punchier'...",
                key="followup", label_visibility="collapsed"
            )
        with c2:
            send = st.button("Send ✦", key="send_btn")

        if send and follow_up:
            if not api_key:
                st.warning("Please enter your OpenAI API key in the sidebar.")
            elif not st.session_state.review_done:
                st.warning("Please wait for the review to complete before sending follow-up questions.")
            elif not st.session_state.openai_history:
                st.warning("No review context found. Please use '🔄 New Review' to start fresh.")
            elif st.session_state.followup_count >= MAX_FOLLOWUPS:
                st.warning(
                    f"You've reached the {MAX_FOLLOWUPS} follow-up limit for this session. "
                    "Use '🔄 New Review' in the sidebar to start fresh."
                )
            else:
                st.session_state.messages.append({"role": "user", "content": follow_up})
                with st.spinner("Krutika AI is thinking..."):
                    try:
                        reply = run_followup(api_key, st.session_state.openai_history, follow_up)
                        st.session_state.openai_history += [
                            {"role": "user", "content": follow_up},
                            {"role": "assistant", "content": reply},
                        ]
                        st.session_state.messages.append({"role": "ai", "content": reply})
                        st.session_state.followup_count += 1
                    except Exception as e:
                        st.session_state.messages.append({"role": "ai", "content": f"❌ Error: {e}"})
                st.rerun()
